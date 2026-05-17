"""Content Parser - Extract text content from various file formats."""

from __future__ import annotations

import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document


@dataclass
class ParsedContent:
    """Result of parsing a file."""
    text: str
    metadata: dict[str, Any]
    chunks: list[str]
    file_type: str


class BaseParser(ABC):
    """Base parser interface."""
    
    @abstractmethod
    def parse(self, content: bytes, filename: str) -> ParsedContent:
        """Parse file content and return extracted text."""
        pass
    
    @abstractmethod
    def supports(self, mime_type: str, filename: str) -> bool:
        """Check if this parser supports the given file type."""
        pass


class TextParser(BaseParser):
    """Parser for plain text files."""

    SUPPORTED_TYPES = {
        "text/plain",
        "text/markdown",
        "text/csv",
        "application/json",
    }
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".yaml", ".yml"}

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text = content.decode("utf-8", errors="replace")
        return ParsedContent(
            text=text,
            metadata={"filename": filename, "size": len(content)},
            chunks=self._chunk(text),
            file_type="text",
        )

    def supports(self, mime_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return mime_type in self.SUPPORTED_TYPES or ext in self.SUPPORTED_EXTENSIONS

    @staticmethod
    def _chunk(text: str, chunk_size: int = 1000, overlap: int = 100) -> list[str]:
        """Split text into overlapping chunks for vector indexing.

        chunk_size is intentionally smaller than the LLM extraction segment
        (8 000 chars) so that each retrieved snippet is focused and precise.
        One LLM segment ≈ 8 vector chunks, guaranteeing full coverage.
        """
        if len(text) <= chunk_size:
            return [text] if text.strip() else []

        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            if chunk.strip():
                chunks.append(chunk)
            start += chunk_size - overlap
        return chunks


class PDFParser(BaseParser):
    """Parser for PDF files using pdfplumber."""
    
    SUPPORTED_TYPES = {"application/pdf"}
    SUPPORTED_EXTENSIONS = {".pdf"}
    
    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text_parts = []
        metadata = {"filename": filename, "pages": 0}
        
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
            tmp.write(content)
            tmp.flush()
            
            with pdfplumber.open(tmp.name) as pdf:
                metadata["pages"] = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
        
        full_text = "\n\n".join(text_parts)
        return ParsedContent(
            text=full_text,
            metadata=metadata,
            chunks=self._chunk(full_text),
            file_type="pdf",
        )
    
    def supports(self, mime_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return mime_type in self.SUPPORTED_TYPES or ext in self.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def _chunk(text: str, chunk_size: int = 1000, overlap: int = 100) -> list[str]:
        """Split text into overlapping chunks for vector indexing.

        chunk_size is intentionally smaller than the LLM extraction segment
        (8 000 chars) so that each retrieved snippet is focused and precise.
        One LLM segment ≈ 8 vector chunks, guaranteeing full coverage.
        Attempts to break at paragraph boundaries to preserve sentence context.
        """
        if len(text) <= chunk_size:
            return [text] if text.strip() else []

        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            # Try to break at paragraph boundary
            chunk = text[start:end]
            last_para = chunk.rfind("\n\n")
            if last_para > chunk_size // 2 and end < len(text):
                end = start + last_para + 2
                chunk = text[start:end]

            if chunk.strip():
                chunks.append(chunk)
            start = end
            if start < len(text) and start > 0:
                start = max(0, start - overlap)
                if start in [s for c in chunks for s in [len(c)]]:
                    start = end  # Avoid infinite loop
        return chunks


class DocxParser(BaseParser):
    """Parser for Word documents."""
    
    SUPPORTED_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    SUPPORTED_EXTENSIONS = {".docx"}
    
    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text_parts = []
        metadata = {"filename": filename, "paragraphs": 0}
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
            tmp.write(content)
            tmp.flush()
            
            doc = Document(tmp.name)
            metadata["paragraphs"] = len(doc.paragraphs)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
        
        full_text = "\n\n".join(text_parts)
        return ParsedContent(
            text=full_text,
            metadata=metadata,
            chunks=TextParser._chunk(full_text),
            file_type="docx",
        )
    
    def supports(self, mime_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return mime_type in self.SUPPORTED_TYPES or ext in self.SUPPORTED_EXTENSIONS


class ExcelParser(BaseParser):
    """Parser for Excel files (.xlsx / .xls / .xlsm).

    Instead of dumping raw cell values to the LLM, this parser produces a
    *structural summary*: sheet names, column headers, row counts, and a
    small sample of values.  The resulting text is short enough for a single
    LLM call and tells the model everything it needs to classify the dataset
    and extract schema-level entities (Dataset, Field, DataType …).
    """

    SUPPORTED_TYPES: set[str] = {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    }
    SUPPORTED_EXTENSIONS: set[str] = {".xlsx", ".xls", ".xlsm"}

    # Maximum rows to include as sample data per sheet
    _SAMPLE_ROWS = 3
    # Maximum sheets to describe in detail (remaining sheets listed by name only)
    _MAX_DETAIL_SHEETS = 5

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        import io
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sections: list[str] = [f"Excel Workbook: {filename}"]
        total_rows = 0
        total_sheets = len(wb.sheetnames)

        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                sections.append(f"\nSheet {sheet_idx + 1}: \"{sheet_name}\" — empty")
                continue

            header = rows[0]
            data_rows = rows[1:]
            n_data = len(data_rows)
            total_rows += n_data
            col_names = [str(h) if h is not None else f"Col{i+1}" for i, h in enumerate(header)]

            if sheet_idx < self._MAX_DETAIL_SHEETS:
                section = [
                    f"\nSheet {sheet_idx + 1}: \"{sheet_name}\" — {n_data:,} rows × {len(col_names)} columns",
                    f"Columns: {', '.join(col_names)}",
                ]
                # Sample rows
                for r_idx, row in enumerate(data_rows[: self._SAMPLE_ROWS], 1):
                    vals = [str(v) if v is not None else "" for v in row[: len(col_names)]]
                    section.append(f"  Row {r_idx}: {' | '.join(vals)}")
                sections.append("\n".join(section))
            else:
                sections.append(f"\nSheet {sheet_idx + 1}: \"{sheet_name}\" — {n_data:,} rows × {len(col_names)} columns (details omitted)")

        wb.close()
        summary_text = "\n".join(sections)
        metadata = {
            "filename": filename,
            "sheets": total_sheets,
            "total_rows": total_rows,
        }
        # No chunking needed — structural summary is already compact
        return ParsedContent(
            text=summary_text,
            metadata=metadata,
            chunks=[summary_text] if summary_text.strip() else [],
            file_type="tabular_data",
        )

    def supports(self, mime_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return mime_type in self.SUPPORTED_TYPES or ext in self.SUPPORTED_EXTENSIONS


class ContentParserService:
    """Service to parse various file formats."""

    def __init__(self) -> None:
        self.parsers: list[BaseParser] = [
            ExcelParser(),   # Must come before TextParser (xlsx is not plain text)
            PDFParser(),
            DocxParser(),
            TextParser(),    # Text parser as fallback
        ]

    def parse(self, content: bytes, filename: str, mime_type: str = "") -> ParsedContent:
        """Parse file content using appropriate parser."""
        for parser in self.parsers:
            if parser.supports(mime_type, filename):
                return parser.parse(content, filename)

        # Fallback to text parser
        return TextParser().parse(content, filename)

    def get_supported_types(self) -> dict[str, list[str]]:
        """Return supported MIME types and extensions."""
        return {
            "mime_types": list(PDFParser.SUPPORTED_TYPES | DocxParser.SUPPORTED_TYPES | TextParser.SUPPORTED_TYPES | ExcelParser.SUPPORTED_TYPES),
            "extensions": list(PDFParser.SUPPORTED_EXTENSIONS | DocxParser.SUPPORTED_EXTENSIONS | TextParser.SUPPORTED_EXTENSIONS | ExcelParser.SUPPORTED_EXTENSIONS),
        }
